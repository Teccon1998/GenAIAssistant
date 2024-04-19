import os
from typing import Dict, Any
from docx import Document
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from langchain.tools import BaseTool, tool

current_dir = os.getcwd()
ROOT_DIR = os.path.join(current_dir, "rootFileSystem")

# Function to add a resume heading
def add_resume_heading(doc: Document, first_name: str, last_name: str, address: str, city: str, state: str, zip_code: str, email: str, phone: str) -> None:
    """ Use this tool to create a header for a resume document """
    heading = doc.add_heading(level=1)
    run = heading.add_run(first_name + ' ' + last_name)
    run.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(address + '•' + city + ',' + state + ' ' + zip_code + ' • ' + email + ' • ' + phone)

# Function to add a resume section with title
def add_section_with_title(doc: Document, title: str) -> None:
    """ This is a tool used to create a section to a title, You are to generate the title yourself, this is for sections of the document like "Experience" " Originizations" "Skills & Interest" and any other neccesary sections """
    doc.add_paragraph(title, style='Heading2')

# Function to add an experience with bullet points
def add_experience(doc: Document, organization: str, position: str, city_state: str, date: str, details: list[str]) -> None:
    """ THis is used to add an experience with bullet points"""
    doc.add_paragraph(organization, style='Heading3')
    p = doc.add_paragraph()
    p.add_run(position + ' ').bold = True
    p.add_run(city_state + ' ' + date).italic = True
    for detail in details:
        p = doc.add_paragraph()
        p.add_run('• ' + detail)


import os
from typing import Dict, Any
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from langchain.tools import BaseTool

current_dir = os.getcwd()
ROOT_DIR = os.path.join(current_dir, "rootFileSystem")

d
