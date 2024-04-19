import os

current_dir = os.getcwd()
ROOT_DIR = os.path.join(current_dir, "rootFileSystem")
#create a new word document instance
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


"""
JSON INFO: 

{
  "firstName": "John",
  "lastName": "Doe",
  "address": "1234 Main Street",
  "city": "Anytown",
  "state": "Anystate",
  "zip_code": "12345",
  "email": "johndoe@example.com",
  "phone": "555-1234",
  "resume_path": "/path/to/resume.docx",
  "jobs": [
    {
      "organization": "Tech Solutions Inc.",
      "position": "Senior Software Engineer",
      "city_state": "New York, NY",
      "date": {
        "start": "2021-08-01",
        "end": "Present"
      },
      "details": [
        "Led the development of a new feature-rich mobile application that increased customer engagement by 70%.",
        "Collaborated with cross-functional teams to integrate advanced AI-driven analytics into company products.",
        "Managed a team of junior developers, overseeing project progress and providing mentorship."
      ]
    },
    {
      "organization": "Innovative Web Services",
      "position": "Web Developer",
      "city_state": "San Francisco, CA",
      "date": {
        "start": "2019-05-01",
        "end": "2021-07-31"
      },
      "details": [
        "Developed responsive website designs compatible with multiple browsers and devices.",
        "Optimized page load times, achieving a 30% improvement on average.",
        "Implemented security best practices to safeguard user data and company intellectual property."
      ]
    }
  ]
}


"""

# Create a new Document
doc = Document()

# Function to add a resume heading
def add_resume_heading(doc, firstName, lastName, address, city, state, zip_code, email, phone):
    """ Use this tool to create a header for a resume document """

    heading = doc.add_heading(level=1)
    run = heading.add_run(firstName + ' ' + lastName)
    run.bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph(address + '•' +  city + ',' + state + ' ' + zip_code  + ' • ' + email +  ' • ' +  phone)

# Function to add a resume section with title
def add_section_with_title(doc, title):
    """ This is a tool used to create a section to a title, You are to generate the title yourself, this is for sections of the document like "Experience" " Originizations" "Skills & Interest" and any other neccesary sections """
    doc.add_paragraph(title, style='Heading2')

# Function to add an experience with bullet points
def add_experience(doc, organization, position, city_state, date, details):
    """ THis is used to add an experience with bullet points"""
    doc.add_paragraph(organization, style='Heading3')
    p = doc.add_paragraph()
    p.add_run(position + ' ').bold = True
    p.add_run(city_state + ' ' + date).italic = True
    for detail in details:
        p = doc.add_paragraph()
        p.add_run('• ' + detail)

# Resume header
add_resume_heading(doc, 'Firstname Lastname\n')

# Education section
add_section_with_title(doc, 'Education')
doc.add_paragraph('HARVARD UNIVERSITY\nDegree, Concentration. GPA [Note: Optional]\nCambridge, MA\nRelevant Coursework: [Note: Optional. Awards and honors can also be listed here.]\nGraduation Date')

# Experience section
add_section_with_title(doc, 'Experience')
add_experience(
    doc,
    'ORGANIZATION',
    'Position Title',
    'City, State (or Remote)',
    'Month Year – Month Year',
    [
        'Beginning with your most recent position, describe your experience, skills, and resulting outcomes in bullet or paragraph form.',
        'Begin each line with an action verb and include details that will help the reader understand your accomplishments, skills, knowledge, abilities, or achievements.',
        'Quantify where possible.',
        'Do not use personal pronouns; each line should be a phrase rather than a full sentence.'
    ]
)

# Leadership & Activities section
add_section_with_title(doc, 'Leadership & Activities')
add_experience(
    doc,
    'ORGANIZATION',
    'Role',
    'City, State',
    'Month Year – Month Year',
    [
        'This section can be formatted similarly to the Experience section, or you can omit descriptions for activities.',
        'If this section is more relevant to the opportunity you are applying for, consider moving this above your Experience section.'
    ]
)

# Skills & Interests section
add_section_with_title(doc, 'Skills & Interests')
doc.add_paragraph(
    'Technical: List computer software and programming languages\n'
    'Language: List foreign languages and your level of fluency\n'
    'Laboratory: List scientific / research lab techniques or tools [If Applicable]\n'
    'Interests: List activities you enjoy that may spark interview conversation'
)

# Save the document
file_path = 'resume_example.docx'
doc.save(file_path)
